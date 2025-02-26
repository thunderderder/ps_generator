import streamlit as st
import requests
import json
from pathlib import Path
import tempfile
import os
from fastapi import UploadFile
import docx
import pandas as pd
import shutil
import time
import streamlit.components.v1 as components

# API 配置
API_BASE_URL = "http://localhost:8000"

# 温度配置
TEMPERATURE_CONFIGS = {
    "分析阶段": 0.7,  # 分析阶段使用较低的温度以保证准确性
    "生成阶段": 2.0,  # 生成阶段使用较高的温度以增加创造性
}

def make_api_request(url: str, request_data: dict, progress_text) -> dict:
    """发送API请求并处理超时"""
    timeout = 300  # 5分钟超时
    start_time = time.time()
    
    try:
        # 更新进度提示，包含预估时间
        progress_text.text(f"请求处理中... (预计需要 {timeout//60} 分钟以内)")
        
        response = requests.post(
            f"{API_BASE_URL}{url}",
            json=request_data,
            timeout=timeout
        )
        
        if response.status_code != 200:
            raise Exception(f"请求失败: HTTP {response.status_code}")
        
        result = response.json()
        if 'response' not in result:
            raise Exception("响应中没有找到 'response' 字段")
        
        # 计算实际用时
        elapsed_time = time.time() - start_time
        progress_text.text(f"处理完成！用时 {elapsed_time:.1f} 秒")
        
        return result
        
    except requests.exceptions.Timeout:
        elapsed_time = time.time() - start_time
        raise Exception(
            f"处理超时（已等待 {elapsed_time:.1f} 秒）。\n"
            "建议：\n"
            "1. 确保系统资源充足\n"
            "2. 检查网络连接\n"
            "3. 检查服务器状态"
        )
    except Exception as e:
        raise Exception(f"请求失败: {str(e)}")

def process_file(file, is_school_info: bool = False) -> str:
    """处理上传的文件并返回内容"""
    if file is None:
        raise ValueError("No file provided")
        
    try:
        # 获取文件内容
        file_content = file.read()
        file.seek(0)
        
        # 创建临时文件
        suffix = Path(file.name).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        
        try:
            if suffix in ['.xls', '.xlsx', '.csv'] and is_school_info:
                # 处理电子表格
                if suffix == '.csv':
                    df = pd.read_csv(tmp_path, encoding='utf-8')
                else:
                    df = pd.read_excel(tmp_path)
                
                # 找到关键列的索引
                result_dict = {}
                for col_idx in range(len(df.columns)):
                    # 获取列名（第一行的值）
                    col_name = str(df.iloc[0, col_idx]).strip() if pd.notna(df.iloc[0, col_idx]) else ""
                    
                    # 如果找到目标列，获取其值（第二行的值）
                    if col_name in ["在读专业", "本科专业", "当前专业"]:
                        value = str(df.iloc[1, col_idx]).strip() if pd.notna(df.iloc[1, col_idx]) else ""
                        result_dict["在读专业"] = value
                    elif col_name in ["申请专业", "目标专业", "意向专业"]:
                        value = str(df.iloc[1, col_idx]).strip() if pd.notna(df.iloc[1, col_idx]) else ""
                        result_dict["申请专业"] = value
                
                if not result_dict:
                    return {"在读专业": "{undergrad_major}", "申请专业": "{target_major}"}
                
                return result_dict
                
            else:
                # 处理其他类型文件
                return process_other_files(file, tmp_path)
                
        finally:
            # 清理临时文件
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception as e:
                st.error(f"Error cleaning up temporary file: {str(e)}")
                
    except Exception as e:
        st.error(f"Error processing file {file.name}: {str(e)}")
        return {"在读专业": "{undergrad_major}", "申请专业": "{target_major}"}

def process_other_files(file, tmp_path):
    """处理非Excel文件"""
    suffix = Path(file.name).suffix.lower()
    if suffix == '.txt':
        with open(tmp_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif suffix == '.docx':
        doc = docx.Document(tmp_path)
        return "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

def process_stream_response(response):
    """处理流式响应"""
    for line in response.iter_lines():
        if line:
            try:
                # 移除 "data: " 前缀
                if line.startswith(b"data: "):
                    line = line[6:]
                data = json.loads(line)
                
                # 处理会话 ID
                if "session_id" in data:
                    yield f"session_id:{data['session_id']}"
                # 处理文本内容
                elif "text" in data:
                    yield data["text"]
                # 处理错误
                elif "error" in data:
                    raise Exception(data["error"])
            except json.JSONDecodeError:
                # 如果不是 JSON 格式，尝试直接解码
                try:
                    text = line.decode('utf-8')
                    if text.startswith("session_id:"):
                        yield text
                    else:
                        yield text
                except:
                    continue

# 获取当前文件的目录
CURRENT_DIR = Path(__file__).parent
# 获取项目根目录
ROOT_DIR = CURRENT_DIR.parent
# 配置文件路径
CONFIG_DIR = ROOT_DIR / 'config'
CONFIG_FILE = CONFIG_DIR / 'default_prompt.txt'

def get_initial_prompt():
    """获取初始提示词模板"""
    try:
        # 确保 config 目录存在
        CONFIG_DIR.mkdir(exist_ok=True)
        # 如果文件不存在，创建默认模板
        if not CONFIG_FILE.exists():
            default_prompt = """我是本科学{undergrad_major}专业的学生，想要申请{target_major}专业，请帮我写一份personal statement。请基本按照以下结构组织内容，你也可以根据内容进行相应的删减或补充，帮我重新组织经历，语言和结构，帮我拓展内容，逻辑清晰，展现我对目标专业的热情，理解和思考。
1. 开篇：介绍申请动机，可以结合对于专业未来方向的发展和思考，展现对目标专业的理解和热情
2. 结合申请人实际的过往经历来谈和申请项目的匹配度。包括但不限于：学术背景、项目经历、实习/研究、个人特质。请注意，描述需要详略得当
3. 重申申请动机，展望未来规划

要求：
1. 使用流畅的学术英语
2. 确保内容真实，基于上传文件中提供的信息
3. 通过具体例子而不是空泛的描述来展示能力
4. 突出与目标专业的契合度"""
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                f.write(default_prompt)
        
        # 读取配置文件
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        st.error(f"Error loading default prompt template: {str(e)}")
        return ""

def save_default_prompt(prompt):
    """保存默认提示词模板"""
    try:
        CONFIG_DIR.mkdir(exist_ok=True)
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            f.write(prompt)
    except Exception as e:
        st.error(f"Error saving default prompt template: {str(e)}")

def load_default_prompt():
    """加载默认提示词模板"""
    try:
        return get_initial_prompt()
    except Exception as e:
        st.error(f"Error loading default prompt: {str(e)}")
        return get_initial_prompt()

# 设置页面配置
st.set_page_config(
    page_title="Personal Statement Generator",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': None,
        'Report a bug': None,
        'About': None
    }
)

# 添加全局样式
st.markdown("""
    <style>
        /* 全局布局样式 */
        .stApp {
            background-color: #f0f2f6;
            color: #1e1e1e;
        }
        
        /* 侧边栏样式 */
        .chat-sidebar {
            background-color: #1e3d59;  /* 深蓝色背景 */
            color: white;
            padding: 1rem;
            border-radius: 10px;
            height: 100%;
        }
        
        /* 主要内容区域样式 */
        .chat-main {
            padding: 1rem;
        }
        
        /* 步骤标题样式 */
        .step-title {
            color: #1e1e1e;
            font-size: 1.2rem;
            font-weight: 500;
            margin-bottom: 1rem;
            padding: 0.5rem 1rem;
            background-color: #ffffff;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        /* 内容区域样式 */
        .content-area {
            color: #1e1e1e;
            padding: 1rem;
            margin-bottom: 1rem;
            min-height: 100px;
        }
        
        /* 步骤说明文本样式 */
        [data-testid="stText"] {
            color: #ffffff !important;  /* 白色文字 */
            font-weight: 500 !important;
            font-size: 1.1rem !important;
        }
        
        /* 步骤说明容器样式 */
        [data-testid="stText"] > div {
            background-color: transparent !important;
            padding: 0 !important;
            border: none !important;
            margin: 0 !important;
        }
        
        /* 进度条样式 */
        .stProgress > div > div > div {
            background-color: rgba(255, 255, 255, 0.2) !important;  /* 半透明白色背景 */
            border-radius: 10px !important;
            height: 8px !important;
        }
        
        .stProgress > div > div > div > div {
            background-color: #ffffff !important;  /* 白色进度条 */
            border-radius: 10px !important;
            height: 8px !important;
        }
        
        /* 聊天框布局样式 */
        .chat-container {
            display: flex;
            gap: 0.5rem;
            margin: -0.3rem;
            padding: 0.3rem;
        }
        
        /* 文件上传组件样式 */
        [data-testid="stFileUploader"] {
            background-color: #ffffff;
            border: 1px solid #dee2e6;
            border-radius: 6px;
            padding: 0.2rem;
            margin-bottom: 0.1rem;
            min-height: 10px;  /* 减小高度 */
        }
        
        /* 修改上传按钮区域样式 */
        [data-testid="stFileUploader"] > div:first-child {
            min-height: 5px;  /* 减小高度 */
            display: flex;
            align-items: center;
            justify-content: space-between;
        }
        
        /* 标题样式 */
        h1, h2, h3, h4, h5, h6 {
            color: #1e1e1e !important;
            margin-bottom: 0.3rem !important;
        }
        
        /* 消息气泡样式 */
        .message-bubble {
            padding: 0.5rem;
            border-radius: 8px;
            margin-bottom: 0.5rem;
            max-width: 100%;
            background-color: #f5f5f5;
        }
        
        /* 按钮样式 */
        .stButton > button,
        .stDownloadButton > button {
            background-color: #1a73e8 !important;
            color: #ffffff !important;
            border: none !important;
            border-radius: 10px !important;
            padding: 0.5rem 1.5rem !important;
            font-weight: 500 !important;
            transition: all 0.2s ease !important;
            width: 100%;
        }
        
        .stButton > button:hover,
        .stDownloadButton > button:hover {
            background-color: #1557b0 !important;
            box-shadow: 0 2px 4px rgba(0,0,0,0.2) !important;
        }
        
        /* 输入框和文本区域样式 */
        .stTextInput > div > div > input,
        .stSelectbox > div > div > select,
        .stTextArea > div > div > textarea {
            background-color: #ffffff !important;
            color: #1e1e1e !important;
            border: 1px solid #dee2e6 !important;
            border-radius: 10px !important;
            padding: 0.75rem !important;
        }
        
        /* 模型描述容器样式 */
        .model-description {
            background-color: #f8f9fa;
            border-radius: 8px;
            padding: 0.5rem;
            margin: 0.5rem 0;
            border: 1px solid #dee2e6;
        }
        
        /* 成功消息样式 */
        .success-message {
            background-color: #e8f5e9;
            color: #1b5e20;
            border: none;
            border-radius: 8px;
            padding: 0.5rem;
            margin: 0.5rem 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        /* 错误消息样式 */
        .error-message {
            background-color: #ffebee;
            color: #c62828;
            border: none;
            border-radius: 8px;
            padding: 0.5rem;
            margin: 0.5rem 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        /* 文件上传状态容器 */
        .upload-status-container {
            display: flex;
            align-items: center;
            gap: 0.1rem;
            margin-bottom: 0.1rem;
            padding: 0.2rem;
            border-radius: 6px;
            background-color: #f8f9fa;
            border: 1px solid #dee2e6;
        }
        
        .upload-status-container.uploaded {
            /* 移除绿色背景，保持与未上传状态一致 */
            background-color: #f8f9fa;
            border-color: #dee2e6;
        }
        
        /* 文件名样式 */
        .file-name {
            flex-grow: 1;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            font-size: 0.8rem;
            color: #1e1e1e;
        }
        
        /* 状态图标样式 */
        .status-icon {
            font-size: 1rem;
            min-width: 10px;
            text-align: center;
        }
        
        /* 文件类型样式 */
        .file-type {
            font-size: 0.75rem;
            color: #666;
            margin-bottom: 0.1rem;
        }
        
        /* 下载区域样式 */
        .download-container {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 1rem;
            margin: 1rem auto;
            max-width: 600px;  /* 限制容器最大宽度 */
        }
        
        /* 调整下拉框宽度 */
        .format-select {
            width: 100px !important;
        }
        
        /* 调整下载按钮宽度 */
        .download-button {
            width: 100px !important;
        }
        
        /* 确保下拉框和按钮在同一行 */
        .stSelectbox, .stDownloadButton {
            margin-bottom: 0 !important;
            padding-bottom: 0 !important;
        }
        
        /* 下载成功提示样式 */
        .download-success {
            background-color: #4CAF50;
            color: white;
            padding: 0.5rem 1rem;
            border-radius: 4px;
            margin-top: 0.5rem;
            text-align: center;
            font-weight: 500;
            max-width: 600px;
            margin-left: auto;
            margin-right: auto;
        }
        
        /* 输入区域容器样式 */
        .input-container {
            background-color: white;
            border-radius: 24px;
            border: 1px solid #e0e0e0;
            padding: 16px;
            margin-bottom: 16px;
            box-shadow: 0 2px 6px rgba(0,0,0,0.05);
        }
        
        /* 输入框容器样式 */
        .stTextArea > div {
            border-radius: 24px !important;
            border: none !important;
            background-color: white !important;
            margin-bottom: 8px !important;
        }
        
        /* 输入框样式 */
        .stTextArea textarea {
            border: none !important;
            padding: 12px 16px !important;
            min-height: 120px !important;
            font-size: 16px !important;
            line-height: 1.5 !important;
            resize: none !important;
            background-color: transparent !important;
        }
        
        /* 按钮容器样式 */
        .button-container {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-top: 8px;
            padding: 4px;
        }
        
        /* 按钮组样式 */
        .button-group {
            display: flex;
            gap: 8px;
        }
        
        /* 按钮样式 */
        .stButton > button {
            border-radius: 20px !important;
            padding: 4px 16px !important;
            font-size: 14px !important;
            font-weight: 500 !important;
            height: 36px !important;
            transition: all 0.2s !important;
            border: 1px solid #e0e0e0 !important;
            background-color: white !important;
            color: #666 !important;
        }
        
        /* 主按钮样式 */
        .stButton > button.primary {
            background-color: #1a73e8 !important;
            color: white !important;
            border: none !important;
        }
        
        .stButton > button:hover {
            transform: translateY(-1px);
            box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        }
        
        /* 圆形图标按钮 */
        .icon-button {
            width: 36px !important;
            height: 36px !important;
            padding: 0 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            border-radius: 50% !important;
        }
        
        /* 生成按钮样式 */
        .generate-button {
            background-color: #1a73e8 !important;
            color: white !important;
            border: none !important;
            border-radius: 20px !important;
            padding: 8px 24px !important;
            font-size: 16px !important;
            font-weight: 500 !important;
            width: 100% !important;
            margin-top: 8px !important;
        }
    </style>
""", unsafe_allow_html=True)

# 初始化会话状态
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = ""
if 'content_result' not in st.session_state:
    st.session_state.content_result = ""
if 'saved_default_prompt' not in st.session_state:
    st.session_state.saved_default_prompt = load_default_prompt()
if 'prompt_template' not in st.session_state:
    st.session_state.prompt_template = st.session_state.saved_default_prompt
if 'school_info_data' not in st.session_state:
    st.session_state.school_info_data = None
if 'session_id' not in st.session_state:
    st.session_state.session_id = None
    
## st.sidebar 下的内容会被渲染到侧边栏
with st.sidebar:
    st.title('Upload Files')
    st.markdown('---')
    
    # Resume 上传
    st.markdown('<div class="file-type">Resume (DOC/DOCX)</div>', unsafe_allow_html=True)
    resume_file = st.file_uploader(
        "Upload your resume",
        type=['doc', 'docx'],
        #label_visibility="collapsed",
        key="resume_uploader"
    )

    # PS 上传
    st.markdown('<div class="file-type">Personal Statement (DOC/DOCX)</div>', unsafe_allow_html=True)
    ps_file = st.file_uploader(
        "Upload your personal statement draft",
        type=['doc', 'docx'],
        label_visibility="collapsed",
        key="ps_uploader"
    )

    # School Info 上传
    st.markdown('<div class="file-type">School Information (XLS/XLSX/CSV)</div>', unsafe_allow_html=True)
    school_file = st.file_uploader(
        "School Information (XLS/XLSX/CSV/TXT)",
        type=['xls', 'xlsx', 'csv'],
        label_visibility="collapsed",
        key="school_uploader"
    )

# 创建页面布局
#st.markdown('<div class="chat-container">', unsafe_allow_html=True)
# 顶部横跨两列的区域
st.markdown('<div class="top-section">', unsafe_allow_html=True)
# Prompt Template Section
status_col, button_col = st.columns([1, 10])
with status_col:
    if st.session_state.get('show_success', False):
        st.markdown("""
            <div class="success-message">
                Default prompt template updated!
            </div>
            """, unsafe_allow_html=True)
        time.sleep(2)
        st.session_state.show_success = False

# 创建一个容器来组织输入区域
with st.container():
    # 如果已上传学校信息，更新提示词模板
    if school_file is not None:
        try:
            school_info_data = process_file(school_file, is_school_info=True)
            st.session_state.school_info_data = school_info_data
            
            # 提取专业信息
            undergrad_major = school_info_data.get("在读专业", "")
            if not undergrad_major:
                undergrad_major = school_info_data.get("本科专业", "")
            if not undergrad_major:
                undergrad_major = school_info_data.get("当前专业", "{undergrad_major}")
                
            target_major = school_info_data.get("申请专业", "")
            if not target_major:
                target_major = school_info_data.get("目标专业", "")
            if not target_major:
                target_major = school_info_data.get("意向专业", "{target_major}")
            
            # 更新提示词模板
            if undergrad_major and target_major:
                st.session_state.prompt_template = st.session_state.prompt_template.replace("{undergrad_major}", undergrad_major).replace("{target_major}", target_major)
        except Exception as e:
            st.error(f"处理学校信息文件时出错: {str(e)}")
    
    # 提示词输入框
    current_prompt = st.text_area(
        "Prompt Template",
        value=st.session_state.prompt_template,
        height=150,
        placeholder="输入你的提示词...",
        label_visibility="collapsed",
        key="prompt_input"
    )
    st.session_state.prompt_template = current_prompt
    
    # 按钮容器
    st.markdown('<div class="button-container">', unsafe_allow_html=True)
    
    # 左侧按钮组
    left_col, right_col = st.columns([7, 1])
    with left_col:
        st.markdown('<div class="button-group">', unsafe_allow_html=True)
        col1, col2, col3 = st.columns([1, 1, 4])
        with col1:
            st.button("🔄", help="重置为默认提示词", key="reset_prompt_btn", use_container_width=True)
        with col2:
            st.button("📝", help="保存为默认提示词", key="save_default", use_container_width=True)
    
    # 生成按钮
    with right_col:
        st.button("生成", key="generate_ps_btn_top", type="primary", use_container_width=True)
    
    st.markdown('</div></div>', unsafe_allow_html=True)
# 显示提示信息
st.markdown("""
    <div style="color: #5f6368; font-size: 0.9rem; margin-top: 0 rem;">
    💡 提示：<br>
    - 当前的修改仅在本次会话中有效<br>
    - 点击 📝 可以将当前内容保存为新的默认模板，请注意保留{undergrad_major}和{target_major}占位符以便自动填充专业信息<br>
    - 点击 🔄 可以恢复为保存的默认模板
    </div>
""", unsafe_allow_html=True)

# 生成按钮
if st.session_state.get("generate_ps_btn_top", False):
    if not all([resume_file, ps_file, school_file]):
        st.error("请先上传所有必需的文件。")
    else:
        try:
            progress_placeholder = st.empty()
            progress_bar = st.progress(0)
            
            # 第一步：分析材料
            progress_placeholder.text("Step 1/2: Analyzing materials...")
            progress_bar.progress(0)
            
            # 读取文件内容到内存
            resume_content = resume_file.read()
            ps_content = ps_file.read()
            school_content = school_file.read()
            
            # 重置文件指针
            resume_file.seek(0)
            ps_file.seek(0)
            school_file.seek(0)
            
            # 准备文件和参数
            files = {
                'resume': ('resume' + Path(resume_file.name).suffix, resume_content),
                'personal_statement': ('ps' + Path(ps_file.name).suffix, ps_content),
                'school_info': ('school' + Path(school_file.name).suffix, school_content)
            }
            data = {
                'temperature': TEMPERATURE_CONFIGS["分析阶段"]
            }
            
            # 创建一个容器用于显示分析结果
            analysis_container = st.container()
            with analysis_container:
                st.markdown("### Material Analysis")
                analysis_output = st.empty()
            analysis_content = ""
            
            # 调用分析 API（使用流式输出）
            with requests.post(
                f"{API_BASE_URL}/analyze_stream",
                files=files,
                data=data,
                stream=True,
                timeout=300
            ) as response:
                if response.status_code != 200:
                    raise Exception(f"Analysis failed: Server returned status code {response.status_code}")
                
                # 处理流式响应
                for chunk in process_stream_response(response):
                    if not st.session_state.session_id and chunk.startswith("session_id:"):
                        st.session_state.session_id = chunk.split(":", 1)[1].strip()
                        continue
                    analysis_content += chunk
                    analysis_output.markdown(
                        f'<div class="message-bubble">{analysis_content}</div>',
                        unsafe_allow_html=True
                    )
            
            # 保存分析结果
            st.session_state.analysis_result = analysis_content
            progress_bar.progress(50)
            
            # 第二步：生成 PS
            progress_placeholder.text("Step 2/2: Generating Personal Statement...")
            
            # 准备生成请求
            generation_data = {
                'analysis': analysis_content,
                'prompt_template': st.session_state.prompt_template,
                'temperature': TEMPERATURE_CONFIGS["生成阶段"],
                'session_id': st.session_state.session_id
            }
            
            # 创建一个容器用于显示生成的内容
            generation_container = st.container()
            with generation_container:
                st.markdown("### Generated Personal Statement")
                content_output = st.empty()
            generated_content = ""
            
            # 使用流式请求
            with requests.post(
                f"{API_BASE_URL}/generate_ps",
                json=generation_data,
                stream=True,
                timeout=300
            ) as response:
                if response.status_code != 200:
                    raise Exception(f"Generation failed: Server returned status code {response.status_code}")
                
                # 处理流式响应
                for chunk in process_stream_response(response):
                    generated_content += chunk
                    content_output.markdown(
                        f'<div class="message-bubble">{generated_content}</div>',
                        unsafe_allow_html=True
                    )
            
            # 保存生成结果
            st.session_state.content_result = generated_content
            progress_bar.progress(100)
            progress_placeholder.text("Complete!")
            
        except Exception as e:
            st.error(f"Error: {str(e)}")
            if 'progress_placeholder' in locals():
                progress_placeholder.empty()
            if 'progress_bar' in locals():
                progress_bar.empty()

# 显示成功消息（如果有）
#if st.session_state.get('show_success', False):
#    st.markdown("""
#       <div class="success-message">
#           默认提示词模板已更新！
#        </div>
#        """, unsafe_allow_html=True)
#    time.sleep(2)
#    st.session_state.show_success = False



# 处理按钮点击事件
if st.session_state.get("reset_prompt_btn", False):
    st.session_state.prompt_template = st.session_state.saved_default_prompt
    st.rerun()

if st.session_state.get("save_default", False):
    st.session_state.saved_default_prompt = st.session_state.prompt_template
    save_default_prompt(st.session_state.prompt_template)
    st.session_state.show_success = True
    st.rerun()

# 下方分列布局
col1, col2 = st.columns(spec=2)

with col1:
    if 'analysis_result' in st.session_state and st.session_state.analysis_result:
        st.markdown(st.session_state.analysis_result)

with col2:
    if 'content_result' in st.session_state and st.session_state.content_result:
        st.markdown("### Generated Personal Statement")
        st.markdown(
            f'<div class="message-bubble">{st.session_state.content_result}</div>',
            unsafe_allow_html=True
        )
        
        # 添加下载选项
        st.markdown('<div class="download-container">', unsafe_allow_html=True)
        
        # 格式选择下拉框
        format_type = st.selectbox(
            "Select format",
            ["Markdown (.md)", "Word (.docx)"],
            label_visibility="collapsed",
            key="format_select_2"
        )
        
        # 下载按钮
        if format_type == "Markdown (.md)":
            if st.download_button(
                "Download Personal Statement",
                st.session_state.content_result,
                "personal_statement.md",
                "text/markdown",
                key="download_btn_2"
            ):
                st.markdown('<div class="download-success">Downloaded as Markdown successfully!</div>', unsafe_allow_html=True)
        else:  # Word (.docx)
            # Word 格式下载
            doc = docx.Document()
            # 添加标题
            doc.add_heading('Personal Statement', 0)
            # 添加内容
            doc.add_paragraph(st.session_state.content_result)
            
            # 保存到临时文件并立即读取
            docx_bytes = None
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc.save(tmp.name)
                with open(tmp.name, 'rb') as f:
                    docx_bytes = f.read()
            try:
                os.unlink(tmp.name)  # 尝试删除临时文件
            except:
                pass  # 忽略删除失败的错误
                
            if st.download_button(
                "Download Personal Statement",
                docx_bytes,
                "personal_statement.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_btn_2"
            ):
                st.markdown('<div class="download-success">Downloaded as Word successfully!</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)  

st.markdown('</div>', unsafe_allow_html=True) 