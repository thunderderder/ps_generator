import os
import json
import logging
import traceback
from pathlib import Path
from typing import Optional, AsyncGenerator, List
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request, Body
from fastapi.middleware.cors import CORSMiddleware
import httpx
from dotenv import load_dotenv
from docx import Document
from matplotlib.table import Cell
import openpyxl
import pandas as pd
import win32com.client
import pythoncom
import asyncio
import sys
from fastapi.responses import StreamingResponse
import tempfile
import docx
import aiofiles
import time
from openai import AsyncOpenAI, OpenAI
from yaml import serialize

# 配置详细的日志格式
logging.basicConfig(
    level=logging.DEBUG,  # 改为 DEBUG 级别
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


# 获取当前文件的目录
current_dir = Path(__file__).resolve().parent
project_root = current_dir.parent.parent

# 尝试加载不同位置的 .env 文件
env_files = [
    project_root / '.env',
    current_dir.parent / '.env',
    current_dir / '.env',
]

env_loaded = False
for env_file in env_files:
    if env_file.exists():
        logger.info(f"找到 .env 文件：{env_file}")
        load_dotenv(env_file)
        env_loaded = True
        break

if not env_loaded:
    logger.error("未找到 .env 文件，请确保文件存在于以下位置之一：")
    for env_file in env_files:
        logger.error(f"- {env_file}")

# 验证环境变量
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not DEEPSEEK_API_KEY:
    logger.error("DEEPSEEK_API_KEY 未在环境变量中设置")
if not OPENAI_API_KEY:
    logger.error("OPENAI_API_KEY 未在环境变量中设置")

# 创建 OpenAI 客户端
deepseek_client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1"
)

openai_client = AsyncOpenAI(
    api_key=OPENAI_API_KEY
)

# API 配置
API_BASE_URL = os.getenv('API_BASE_URL', 'http://localhost:8000')
logger.info(f"Using API base URL: {API_BASE_URL}")

# 分析阶段 API 配置
"""
ANALYSIS_API_CONFIG = {
    "model": "gpt-4o",  # 使用 GPT-4 with vision
    "max_tokens": 8000,
    "stream": True,
    "client": openai_client,
    "system_prompt": "You are a specialized assistant with expertise in crafting compelling and personalized application materials for master's program applications."
}
"""

# DeepSeek API 配置
ANALYSIS_API_CONFIG = {
    "model": "deepseek-reasoner",  # 必须完全匹配，区分大小写
    "max_tokens": 8000,  # 最大回答长度
    "stream": True,
    "client": deepseek_client,
    "temperature": 1.0,
    "system_prompt": "You are a specialized assistant with expertise in crafting compelling and personalized application materials for master's program applications."
}

# 生成阶段 API 配置
GENERATION_API_CONFIG = {
    "model": "gpt-4o",
    "max_tokens": 8000,
    "stream": True,
    "client": openai_client,
    "system_prompt": """
You are an expert assistant specializing in crafting impactful and tailored application materials 
for master's program applications. Your primary goal is to help users create a polished, persuasive, 
and passion-driven Personal Statement that resonates with admissions committees and professors. 
You possess exceptional skills in academic English, logical analysis, and professional expression, 
and have a strong understanding of the stylistic and substantive preferences of institutions across 
different regions. You are also highly familiar with the conventions of Personal Statement writing 
and can adapt content to suit diverse academic and professional backgrounds.

When generating a Personal Statement, adhere to the following principles:

1. Use Fluent Academic English: 
Ensure the language is polished, professional, and appropriate for an academic audience. 
Avoid overly casual expressions or vague statements.

2. Base Content on Provided Information: 
Ensure the statement is authentic and accurately reflects the applicant's experiences and achievements, 
strictly using the information provided in the uploaded materials.

3. Highlight Skills Through Specific Examples: 
Avoid generic descriptions of abilities. Instead, use detailed and concrete examples 
(e.g., specific projects, internships, or research) to demonstrate the applicant's skills and achievements.

4. Emphasize the Fit with the Target Program: 
Clearly showcase the alignment between the applicant's background, experiences, and the goals of the target program. 
Emphasize how their unique qualifications make them a strong candidate.

The overall structure of the Personal Statement should follow these key sections:

1. Opening Paragraph: 
Introduce the applicant's motivation for pursuing the program. This can include reflections on their 
understanding of the field, its future development, and their enthusiasm for contributing to it.

2. Main Body - Demonstrating Fit: 
Draw from the applicant's academic background, project work, research, internships, or other relevant 
experiences to demonstrate their preparedness for the program. Ensure the description is well-organized 
and prioritizes experiences that showcase a strong connection to the target discipline.

3. Conclusion - Reaffirming Motivation and Future Goals: 
Reiterate the applicant's motivation for applying and outline their long-term career aspirations 
and how the program will help them achieve these goals.

Your task is to organize and refine the provided content, ensuring the statement is well-structured, 
logically compelling, and vividly conveys the applicant's passion, insights, and alignment with the program. 
You may reorganize and enhance the language, clarify ideas, and adjust emphasis to highlight the applicant's 
strongest qualifications. Always maintain a professional and engaging tone, ensuring the statement leaves 
a lasting impression on the reader.

    """
}

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

# 添加一个常量定义最大处理长度
MAX_CHUNK_SIZE = 8000  # 根据实际模型限制调整

# 添加应用状态
class AppState:
    def __init__(self):
        self.temp_files = {}  # 存储会话ID到临时文件路径的映射

app.state.app_state = AppState()

# 清理函数
async def cleanup_temp_files(session_id: str):
    """清理指定会话的临时文件"""
    if session_id in app.state.app_state.temp_files:
        files = app.state.app_state.temp_files[session_id]
        for path in files:
            if path.exists():
                try:
                    path.unlink()
                    logger.info(f"删除临时文件：{path}")
                except Exception as e:
                    logger.error(f"清理临时文件失败 {path}: {str(e)}")
        del app.state.app_state.temp_files[session_id]

# 1. 工具函数
async def call_ai_with_retry_stream(prompt: str, api_config: dict, max_retries: int = 3):
    """调用AI服务生成文本，支持流式输出"""
    retry_count = 0
    last_error = None
    
    while retry_count < max_retries:
        try:
            logger.info("创建流式请求...")
            # 使用 API SDK 进行流式调用
            if api_config["model"] == "deepseek-reasoner":
                # DeepSeek API 不支持 async，使用同步调用
                stream = api_config["client"].chat.completions.create(
                    model=api_config["model"],
                    messages=[
                        {"role": "system", "content": api_config["system_prompt"]},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=api_config["max_tokens"],
                    stream=True,
                    timeout=300  # 增加超时时间到5分钟
                )
                
                content_received = False
                last_content_time = time.time()
                logger.info("开始处理 DeepSeek 流式响应...")
                
                reasoning_content = ""
                content = ""
                
                try:
                    # 处理 DeepSeek 流式响应
                    for chunk in stream:  # 不使用 async for
                        current_time = time.time()
                        
                        if current_time - last_content_time > 180:
                            raise Exception("Connection stalled - no content received for 180 seconds")
                        
                        if chunk.choices[0].delta.reasoning_content:
                            content_received = True
                            last_content_time = current_time
                            reasoning_content += chunk.choices[0].delta.reasoning_content
                            logger.debug(f"收到推理内容片段: {chunk.choices[0].delta.reasoning_content[:50]}...")
                            yield chunk.choices[0].delta.reasoning_content
                        elif hasattr(chunk.choices[0].delta, 'content') and chunk.choices[0].delta.content:
                            content_received = True
                            last_content_time = current_time
                            content += chunk.choices[0].delta.content
                            logger.debug(f"收到回答内容片段: {chunk.choices[0].delta.content[:50]}...")
                            yield chunk.choices[0].delta.content
                    
                    if not content_received:
                        raise Exception("No content received from the API")
                    
                    logger.info("DeepSeek 流式响应处理完成")
                    return
                    
                except Exception as e:
                    logging.error(f"DeepSeek Streaming error: {str(e)}")
                    raise e
            else:
                # 其他 API 使用异步调用
                stream = await api_config["client"].chat.completions.create(
                    model=api_config["model"],
                    messages=[
                        {"role": "system", "content": api_config["system_prompt"]},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=api_config["max_tokens"],
                    stream=True,
                    timeout=300  # 增加超时时间到5分钟
                )
                
                content_received = False
                last_content_time = time.time()
                logger.info("开始处理流式响应...")
                
                try:
                    # 处理流式响应
                    async for chunk in stream:
                        current_time = time.time()
                        
                        if current_time - last_content_time > 180:
                            raise Exception("Connection stalled - no content received for 180 seconds")
                        
                        if hasattr(chunk.choices[0].delta, 'content') and chunk.choices[0].delta.content is not None:
                            content_received = True
                            last_content_time = current_time
                            logger.debug(f"收到内容片段: {chunk.choices[0].delta.content[:50]}...")
                            yield chunk.choices[0].delta.content
                    
                    if not content_received:
                        raise Exception("No content received from the API")
                    
                    logger.info("流式响应处理完成")
                    return
                    
                except Exception as e:
                    logging.error(f"Streaming error: {str(e)}")
                    raise e
                
        except Exception as e:
            retry_count += 1
            last_error = e
            logging.error(f"Attempt {retry_count} failed: {str(e)}")
            if retry_count < max_retries:
                wait_time = 2 ** retry_count  # 指数退避
                logger.info(f"等待 {wait_time} 秒后重试...")
                await asyncio.sleep(wait_time)
                logger.info("重新初始化客户端...")
                # 重新初始化客户端
                if api_config["model"] == "deepseek-reasoner":
                    api_config["client"] = OpenAI(
                        api_key=os.getenv("DEEPSEEK_API_KEY"),
                        base_url="https://api.deepseek.com/v1"
                    )
                else:
                    api_config["client"] = AsyncOpenAI(
                        api_key=os.getenv("OPENAI_API_KEY")
                    )
                continue
            break
            
    raise Exception(f"在 {max_retries} 次尝试后调用仍然失败：{str(last_error)}")

def convert_doc_to_docx(doc_path):
    """转换 doc 文件为 docx 格式"""
    try:
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False  # 禁用警告
        except Exception as e:
            raise Exception("无法启动 Word，请确保已安装 Microsoft Word 且以管理员权限运行程序。错误信息：" + str(e))
        
        try:
            # 创建一个新的空白文档
            new_doc = word.Documents.Add()
            
            # 读取原文档内容
            with open(doc_path, 'rb') as file:
                content = file.read()
            
            # 创建临时文件
            temp_path = os.path.join(os.path.dirname(doc_path), "temp_" + os.path.basename(doc_path))
            with open(temp_path, 'wb') as file:
                file.write(content)
            
            try:
                # 尝试打开临时文件
                doc = word.Documents.Open(os.path.abspath(temp_path))
                docx_path = doc_path + "x"  # Add 'x' to make it .docx
                doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)  # 16 = docx format
                doc.Close()
                return docx_path
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    try:
                        os.remove(temp_path)
                    except:
                        pass
        except Exception as e:
            if "检测到此文件存在一个问题" in str(e):
                raise Exception("Word 安全设置阻止打开文件。请手动打开文件，点击'启用编辑'，然后重新上传，或者将文件另存为 .docx 格式后重试。")
            raise Exception("转换文档时出错，请确保文档未被其他程序占用。错误信息：" + str(e))
        finally:
            try:
                word.Quit()
            except:
                pass
    except Exception as e:
        raise Exception(f"文档转换失败: {str(e)}")
    finally:
        pythoncom.CoUninitialize()

def read_document(file_path):
    """读取文档内容"""
    try:
        logger.info(f"开始读取文档: {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if not os.path.exists(file_path):
            raise Exception(f"文件不存在: {file_path}")
            
        logger.info(f"文件格式: {file_ext}")
        
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                logger.info("成功读取文本文件")
            elif file_ext == '.doc':
                # 将 .doc 转换为 .docx
                docx_path = convert_doc_to_docx(file_path)
                doc = Document(docx_path)
                
                # 读取段落文本
                paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                
                # 读取表格内容
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:  # 只添加非空行
                            tables_text.append(" | ".join(row_text))
                
                # 合并段落和表格内容
                text = "\n\n".join(paragraphs_text)
                if tables_text:
                    text += "\n\n表格内容：\n" + "\n".join(tables_text)
                
                # 清理临时文件
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                logger.info("成功读取 DOC 文件")
            elif file_ext == '.docx':
                doc = Document(file_path)
                
                # 读取段落文本
                paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                
                # 读取表格内容
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:  # 只添加非空行
                            tables_text.append(" | ".join(row_text))
                
                # 合并段落和表格内容
                text = "\n\n".join(paragraphs_text)
                if tables_text:
                    text += "\n\n表格内容：\n" + "\n".join(tables_text)
                    
                logger.info("成功读取 DOCX 文件")
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            logger.info(f"文档内容长度: {len(text)} 字符")
            if tables_text:
                logger.info(f"包含表格行数: {len(tables_text)}")
            return text
        except Exception as e:
            logger.error(f"读取文档失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"读取文档失败: {str(e)}")
            
    except Exception as e:
        logger.error(f"处理文档时出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"处理文档时出错: {str(e)}")

def read_school_info(file_path):
    """读取学校信息"""
    try:
        logger.info(f"开始读取学校信息: {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if not os.path.exists(file_path):
            raise Exception(f"文件不存在: {file_path}")
            
        logger.info(f"文件格式: {file_ext}")
        
        try:
            if file_ext == '.txt':
                # 处理文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                data = [line.strip().split(',') for line in content.splitlines() if line.strip()]
                formatted_data = {}
                for line in data:
                    if len(line) >= 2:
                        formatted_data[line[0].strip()] = line[1].strip()
            elif file_ext in ['.csv', '.xls', '.xlsx']:
                # 读取 Excel/CSV 文件
                if file_ext == '.csv':
                    df = pd.read_csv(file_path, encoding='utf-8')
                else:
                    df = pd.read_excel(file_path)
                
                logger.info(f"原始数据形状: {df.shape}")
                logger.info(f"列名: {df.columns.tolist()}")
                
                # 初始化结果字典
                formatted_data = {}
                
                # 遍历所有单元格查找关键信息
                for row in range(df.shape[0]):
                    for col in range(df.shape[1]):
                        cell_value = str(df.iloc[row, col]).strip() if pd.notna(df.iloc[row, col]) else ""
                        
                        # 检查是否是关键字段
                        if cell_value in ["在读专业", "本科专业", "当前专业", "申请专业", "目标专业", "意向专业"]:
                            # 获取下一行的值（如果存在）
                            if row + 1 < df.shape[0]:
                                next_value = str(df.iloc[row + 1, col]).strip() if pd.notna(df.iloc[row + 1, col]) else ""
                                if next_value:  # 只保存非空值
                                    if cell_value in ["在读专业", "本科专业", "当前专业"]:
                                        formatted_data["在读专业"] = next_value
                                    elif cell_value in ["申请专业", "目标专业", "意向专业"]:
                                        formatted_data["申请专业"] = next_value
                
                # 如果没有找到任何信息，使用默认值
                if not formatted_data:
                    formatted_data = {
                        "在读专业": "A专业",
                        "申请专业": "B专业"
                    }
                
                logger.info(f"处理后的数据条数: {len(formatted_data)}")
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            logger.info(f"最终数据示例: {formatted_data}")
            return formatted_data
            
        except Exception as e:
            logger.error(f"读取表格失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"读取表格失败: {str(e)}")
            
    except Exception as e:
        logger.error(f"处理表格时出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise Exception(f"处理表格时出错: {str(e)}")

async def write_bytes_to_file(path: Path, content: bytes):
    """异步写入文件内容"""
    async with aiofiles.open(path, 'wb') as f:
        await f.write(content)

def extract_majors_from_school_info(school_info_data: dict) -> tuple[str, str]:
    """从学校信息中提取专业信息"""
    try:
        # 直接查找对应的列名和值
        undergrad_major = school_info_data.get("在读专业", "")
        if not undergrad_major:
            undergrad_major = school_info_data.get("本科专业", "")
        if not undergrad_major:
            undergrad_major = school_info_data.get("当前专业", "A专业")
            
        target_major = school_info_data.get("申请专业", "")
        if not target_major:
            target_major = school_info_data.get("目标专业", "")
        if not target_major:
            target_major = school_info_data.get("意向专业", "B专业")
        
        logger.info(f"提取到的专业信息 - 在读专业: {undergrad_major}, 申请专业: {target_major}")
        return (undergrad_major, target_major)
        
    except Exception as e:
        logger.error(f"提取专业信息时出错: {str(e)}")
        return ("A专业", "B专业")  # 使用默认值

def split_text_into_chunks(text: str, max_chunk_size: int = 4000) -> list[str]:
    """将长文本分割成较小的块，确保每个块的大小不超过max_chunk_size"""
    # 按段落分割
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = []
    current_size = 0
    
    for paragraph in paragraphs:
        paragraph_size = len(paragraph)
        if current_size + paragraph_size > max_chunk_size and current_chunk:
            # 当前块已满，保存并开始新块
            chunks.append('\n\n'.join(current_chunk))
            current_chunk = [paragraph]
            current_size = paragraph_size
        else:
            # 添加到当前块
            current_chunk.append(paragraph)
            current_size += paragraph_size
    
    # 添加最后一个块
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks

# 2. 核心业务逻辑
async def analyze_text_chunks(chunks: List[str], model: str, temperature: float) -> str:
    """分析文本块并生成综合分析结果"""
    analysis_results = []
    
    for i, chunk in enumerate(chunks):
        logger.info(f"处理第 {i+1}/{len(chunks)} 个文本块")
        chunk_prompt = f"请分析以下文本并提供见解：\n\n{chunk}"
        
        try:
            result = await call_ai_with_retry_stream(chunk_prompt, api_config=ANALYSIS_API_CONFIG)
            if isinstance(result, dict) and 'content' in result:
                analysis_results.append(result['content'])
            elif isinstance(result, str):
                analysis_results.append(result)
            else:
                logger.error(f"意外的响应格式：{result}")
                raise ValueError(f"AI 服务返回了意外的响应格式：{result}")
                
        except Exception as e:
            logger.error(f"处理文本块 {i+1} 失败：{str(e)}")
            logger.error(traceback.format_exc())
            # 不立即失败，记录错误并继续处理其他块
            analysis_results.append(f"[处理此部分时出错: {str(e)}]")
            continue
    
    if not analysis_results:
        raise HTTPException(
            status_code=503,
            detail="所有文本块处理均失败"
        )
    
    # 合并所有分析结果
    try:
        combined_analysis = '\n\n---\n\n'.join(
            result for result in analysis_results 
            if not result.startswith('[处理此部分时出错')
        )
        
        if not combined_analysis.strip():
            raise HTTPException(
                status_code=503,
                detail="生成的内容为空"
            )
            
        logger.info("成功合并所有分析结果")
        return combined_analysis
        
    except Exception as e:
        logger.error(f"合并分析结果时出错：{str(e)}")
        logger.error(f"分析结果：{analysis_results}")
        raise HTTPException(
            status_code=500,
            detail=f"合并分析结果失败：{str(e)}"
        )

async def process_materials(resume_text: str, ps_text: str, school_info_data: dict, model: str, temperature: float, custom_prompt: str = None) -> str:
    """统一的材料处理函数 - 分析阶段"""
    
    # 标准化模型名称
    model = model.lower().strip()
    
    # 从学校信息中提取专业信息
    undergrad_major, target_major = extract_majors_from_school_info(school_info_data)
    
    # 合并所有材料
    combined_text = f"""简历信息：
{resume_text}

个人陈述调查表信息：
{ps_text}

申请学校信息：
{json.dumps(school_info_data, ensure_ascii=False, indent=2)}"""

    # 构建分析提示词
    analysis_prompt = f"""请分析以下申请材料，提取关键信息和亮点，并标注信息的来源（如：简历调查表、个人陈述调查表、申请学校信息表等）为生成个人陈述做准备。
请重点关注以下几个方面和目标专业的匹配度：
1. 申请人的学术背景和专业知识与目标专业的匹配度
2. 相关的项目经历和实习经验与目标专业的匹配度
3. 研究经历和成果与目标专业的匹配度

申请人背景：
- 本科专业：{undergrad_major}
- 目标专业：{target_major}

{combined_text}

请提供详细的分析，包括：
1. 申请人的主要优势和特点
2. 可以重点突出的经历和成果
3. 需要特别说明或解释的内容
4. 个人陈述写作的整体思路建议"""

    # 判断是否需要分块处理
    if len(combined_text) > MAX_CHUNK_SIZE:
        logger.info(f"文本长度({len(combined_text)})超过限制，使用分块处理")
        # 分块处理
        chunks = split_text_into_chunks(combined_text)
        return await analyze_text_chunks(chunks, model, temperature)
    else:
        logger.info(f"文本长度({len(combined_text)})在限制内，使用整体处理")
        # 使用流式输出进行分析
        full_response = ""
        async for chunk in call_ai_with_retry_stream(analysis_prompt, api_config=ANALYSIS_API_CONFIG):
            full_response += chunk
        return full_response

# 3. API 端点
@app.post("/analyze_stream")
async def analyze_materials_stream(
    resume: UploadFile = File(...),
    personal_statement: UploadFile = File(...),
    school_info: UploadFile = File(...),
    prompt_template: str = Form(None),
    session_id: str = Form(None)  # Make session_id optional
):
    """流式分析材料端点"""
    try:
        logger.info("开始处理流式分析请求")
        logger.info(f"接收到的文件：resume={resume.filename}, ps={personal_statement.filename}, school={school_info.filename}")
        
        # Generate session_id if not provided
        if not session_id:
            session_id = str(int(time.time() * 1000))  # Use timestamp as session ID
            logger.info(f"自动生成会话ID: {session_id}")
        
        # 读取所有文件内容到内存
        logger.info("读取文件内容到内存")
        resume_content = await resume.read()
        ps_content = await personal_statement.read()
        school_info_content = await school_info.read()
        
        logger.info(f"文件大小：resume={len(resume_content)}字节, ps={len(ps_content)}字节, school={len(school_info_content)}字节")
        
        # 重置文件指针，以防后续还需要读取
        await resume.seek(0)
        await personal_statement.seek(0)
        await school_info.seek(0)
        
        async def generate():
            try:
                logger.info("开始生成响应")
                # 创建临时文件并写入内容
                resume_path = Path(tempfile.gettempdir()) / f"resume_{resume.filename}"
                ps_path = Path(tempfile.gettempdir()) / f"ps_{personal_statement.filename}"
                school_info_path = Path(tempfile.gettempdir()) / f"school_{school_info.filename}"
                
                # 保存临时文件路径到应用状态
                app.state.app_state.temp_files[session_id] = [resume_path, ps_path, school_info_path]
                
                try:
                    # 写入临时文件
                    logger.info("写入临时文件")
                    await write_bytes_to_file(resume_path, resume_content)
                    await write_bytes_to_file(ps_path, ps_content)
                    await write_bytes_to_file(school_info_path, school_info_content)
                    
                    # 处理文件内容
                    logger.info("开始处理文件内容")
                    resume_text = read_document(str(resume_path))
                    logger.info("简历处理完成")
                    ps_text = read_document(str(ps_path))
                    logger.info("个人陈述处理完成")
                    school_info_data = read_school_info(str(school_info_path))
                    logger.info("学校信息处理完成")
                    
                    # 从学校信息中提取专业信息
                    undergrad_major, target_major = extract_majors_from_school_info(school_info_data)
                    logger.info(f"提取到的专业信息：本科={undergrad_major}, 目标={target_major}")
                    
                    # 构建分析提示词
                    analysis_prompt = f"""请分析以下申请材料，提取关键信息和亮点为撰写PS提供思路，每个关键点都需要标注信息的来源（如：简历调查表、个人陈述调查表、申请学校信息表等）为生成个人陈述做准备。
请重点关注以下几个方面和目标专业的匹配度：
1. 申请人的学术背景和专业知识与目标专业的匹配度
2. 相关的项目经历和实习经验与目标专业的匹配度
3. 研究经历和成果与目标专业的匹配度

特别注意，如果申请人背景专业和目标专业不一致，要仔细思考本科专业和目标专业的关联是什么，怎么通过过往的经历串在一起。
- 本科专业：{undergrad_major}
- 目标专业：{target_major}

简历信息：
{resume_text}

个人陈述调查表信息：
{ps_text}

申请学校信息：
{json.dumps(school_info_data, ensure_ascii=False, indent=2)}

请提供详细的分析，包括：
1. 申请人的主要优势和特点
2. 可以重点突出的经历和成果
3. 需要特别说明或解释的内容
4. 个人陈述写作的整体思路建议"""

                    logger.info("开始调用AI服务进行分析")
                    content_generated = False
                    
                    # First yield the session ID
                    yield f"data: {json.dumps({'session_id': session_id}, ensure_ascii=False)}\n\n".encode('utf-8')
                    
                    # 使用流式输出进行分析
                    async for chunk in call_ai_with_retry_stream(analysis_prompt, api_config=ANALYSIS_API_CONFIG):
                        if chunk:
                            content_generated = True
                            logger.debug(f"生成内容片段：{chunk[:50]}...")
                            yield f"data: {json.dumps({'text': chunk}, ensure_ascii=False)}\n\n".encode('utf-8')
                    
                    if not content_generated:
                        error_msg = "分析过程未产生任何内容"
                        logger.error(error_msg)
                        yield f"data: {json.dumps({'error': error_msg}, ensure_ascii=False)}\n\n".encode('utf-8')
                        
                except Exception as e:
                    error_msg = f"分析过程发生错误: {str(e)}"
                    logger.error(error_msg)
                    logger.error(traceback.format_exc())
                    yield f"data: {json.dumps({'error': error_msg}, ensure_ascii=False)}\n\n".encode('utf-8')
                
            except Exception as e:
                error_msg = f"分析过程发生错误: {str(e)}"
                logger.error(error_msg)
                logger.error(traceback.format_exc())
                yield f"data: {json.dumps({'error': error_msg}, ensure_ascii=False)}\n\n".encode('utf-8')

        return StreamingResponse(
            generate(),
            media_type="text/event-stream"
        )
        
    except Exception as e:
        logger.error(f"文件读取失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate_ps")
async def generate_ps_endpoint(
    request: Request
):
    """生成个人陈述端点"""
    try:
        # 获取请求数据
        data = await request.json()
        analysis = data.get("analysis")
        prompt_template = data.get("prompt_template")
        session_id = data.get("session_id")
        temperature = data.get("temperature", 2)
        
        logger.info(f"收到生成请求 - session_id: {session_id}")
        logger.info(f"分析结果长度: {len(analysis) if analysis else 0}")
        logger.info(f"提示词模板长度: {len(prompt_template) if prompt_template else 0}")
        
        if not all([analysis, prompt_template, session_id]):
            missing = []
            if not analysis:
                missing.append("analysis")
            if not prompt_template:
                missing.append("prompt_template")
            if not session_id:
                missing.append("session_id")
            raise HTTPException(
                status_code=400, 
                detail=f"Missing required parameters: {', '.join(missing)}"
            )
            
        if session_id not in app.state.app_state.temp_files:
            logger.error(f"会话 ID {session_id} 不存在")
            logger.info(f"当前可用的会话 ID: {list(app.state.app_state.temp_files.keys())}")
            raise HTTPException(
                status_code=400, 
                detail=f"Session expired or invalid: {session_id}"
            )
            
        # 获取临时文件路径
        temp_files = app.state.app_state.temp_files[session_id]
        if len(temp_files) != 3:
            logger.error(f"临时文件数量不正确: {len(temp_files)}")
            raise HTTPException(
                status_code=500, 
                detail="Incomplete session data"
            )
            
        resume_path, ps_path, school_info_path = temp_files
        logger.info(f"找到临时文件: {resume_path}, {ps_path}, {school_info_path}")
        
        # 读取文件内容
        try:
            resume_text = read_document(str(resume_path))
            ps_text = read_document(str(ps_path))
            school_info_data = read_school_info(str(school_info_path))
            logger.info("成功读取所有文件内容")
        except Exception as e:
            logger.error(f"读取文件内容失败: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to read session files: {str(e)}"
            )
        
        # 构建生成提示词
        generation_prompt = f"""基于用户上传的材料、材料分析结果和PS撰写要求，生成一份完整的个人陈述。

用户上传的材料：
简历信息：
{resume_text}

个人陈述调查表信息：
{ps_text}

申请学校信息：
{json.dumps(school_info_data, ensure_ascii=False, indent=2)}

分析结果（包括申请人的背景、优势、经历等）：
{analysis}

用户要求：
{prompt_template}

请按照以上要求，生成一份完整的个人陈述。要求：
1. 使用流畅的学术英语
2. 确保内容真实，基于分析结果中提供的信息
3. 通过具体例子展示申请人的能力和特点
4. 突出申请人与目标专业的契合度
5. 保持逻辑清晰，段落衔接自然"""

        # 记录提示词长度
        prompt_length = len(generation_prompt)
        logger.info(f"生成提示词总长度：{prompt_length} 字符")
        logger.debug(f"分析结果长度：{len(analysis)} 字符")
        logger.debug(f"模板长度：{len(prompt_template)} 字符")

        async def generate():
            try:
                content_generated = False
                error_occurred = False
                
                # 使用生成阶段的 API 配置
                async for chunk in call_ai_with_retry_stream(generation_prompt, api_config=GENERATION_API_CONFIG):
                    if chunk:
                        content_generated = True
                        logger.debug(f"生成内容片段：{chunk[:50]}...")
                        # 确保正确的编码
                        json_str = json.dumps({'text': chunk}, ensure_ascii=False)
                        yield f"data: {json_str}\n\n".encode('utf-8')
                
                if not content_generated and not error_occurred:
                    error_msg = "生成过程未产生任何内容，请检查API密钥和网络连接"
                    logger.error(error_msg)
                    json_str = json.dumps({'error': error_msg}, ensure_ascii=False)
                    yield f"data: {json_str}\n\n".encode('utf-8')
                    
            except Exception as e:
                error_occurred = True
                error_msg = f"生成失败: {str(e)}"
                logger.error(error_msg)
                logger.error(traceback.format_exc())
                json_str = json.dumps({'error': error_msg}, ensure_ascii=False)
                yield f"data: {json_str}\n\n".encode('utf-8')

        return StreamingResponse(
            generate(),
            media_type="text/event-stream; charset=utf-8"
        )
        
    except Exception as e:
        logger.error(f"生成请求失败: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload")
async def upload_files(
    resume: UploadFile = File(...),
    personal_statement: UploadFile = File(...),
    school_info: UploadFile = File(...),
    model: str = Form("deepseek-api"),
    temperature: float = Form(0.7),
    prompt_template: str = Form(None)
):
    try:
        logger.info(f"Received upload request - Files: {resume.filename}, {personal_statement.filename}, {school_info.filename}")
        logger.info(f"Model: {model}, Temperature: {temperature}")
        
        # 初始化上传文件列表用于跟踪清理
        uploaded_files = []
        
        ext = os.path.splitext(school_info.filename)[1].lower()
        if ext not in ['.xls', '.xlsx', '.csv', '.txt']:
            raise Exception(f"文件 {school_info.filename} 格式不正确。请上传 .xls、.xlsx、.csv 或 .txt 格式的文件。")

        # 创建上传目录
        if not os.path.exists(UPLOAD_DIR):
            try:
                os.makedirs(UPLOAD_DIR)
                logger.info(f"创建上传目录: {UPLOAD_DIR}")
            except Exception as e:
                logger.error(f"创建上传目录失败: {str(e)}")
                raise Exception(f"创建上传目录失败: {str(e)}")

        resume_path = os.path.join(UPLOAD_DIR, resume.filename)
        ps_path = os.path.join(UPLOAD_DIR, personal_statement.filename)
        school_path = os.path.join(UPLOAD_DIR, school_info.filename)
        
        logger.info("开始保存上传的文件")
        # Save uploaded files
        try:
            with open(resume_path, "wb") as f:
                content = await resume.read()
                logger.info(f"读取简历文件，大小: {len(content)} 字节")
                f.write(content)
            uploaded_files.append(resume_path)
            
            with open(ps_path, "wb") as f:
                content = await personal_statement.read()
                logger.info(f"读取个人陈述文件，大小: {len(content)} 字节")
                f.write(content)
            uploaded_files.append(ps_path)
            
            with open(school_path, "wb") as f:
                content = await school_info.read()
                logger.info(f"读取学校信息文件，大小: {len(content)} 字节")
                f.write(content)
            uploaded_files.append(school_path)
            
            logger.info("所有文件保存成功")
            
        except Exception as e:
            logger.error(f"保存上传文件失败: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"保存上传文件失败: {str(e)}")
        
        try:
            # Read documents
            logger.info("开始读取文档内容")
            resume_text = read_document(resume_path)
            logger.info("简历读取完成")
            
            ps_text = read_document(ps_path)
            logger.info("个人陈述读取完成")
            
            # 读取学校信息
            school_info_data = read_school_info(school_path)
            logger.info("学校信息读取完成")
            
            # 使用统一的处理函数
            generated_text = await process_materials(
                resume_text,
                ps_text,
                school_info_data,
                model,
                temperature,
                prompt_template
            )
            
            return {"status": "success", "analysis": generated_text}
            
        except Exception as e:
            logger.error(f"处理文件内容时出错: {str(e)}")
            logger.error(traceback.format_exc())
            raise Exception(f"处理文件内容时出错: {str(e)}")
            
    except Exception as e:
        logger.error(f"上传处理失败: {str(e)}")
        logger.error(traceback.format_exc())
        return {"status": "error", "message": str(e)}
    finally:
        # Clean up uploaded files
        logger.info("清理上传的文件")
        for file_path in uploaded_files:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"删除文件: {file_path}")
            except Exception as e:
                logger.error(f"删除文件失败: {file_path}, 错误: {str(e)}")
                pass  # Ignore cleanup errors

async def process_file(file: UploadFile, is_school_info: bool = False) -> str:
    """处理上传的文件并返回内容"""
    file_path = None
    try:
        # 读取上传的文件内容
        content = await file.read()
        
        # 创建临时文件
        suffix = Path(file.filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            file_path = tmp.name
            # 写入临时文件
            tmp.write(content)
            
        # 根据文件类型处理内容
        if suffix == '.txt':
            # 尝试不同的编码方式读取文本文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1']
            text_content = None
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            if text_content is None:
                raise ValueError(f"Unable to decode {file.filename} with any supported encoding")
            return text_content
            
        elif suffix == '.docx':
            # 处理 Word 文档
            doc = docx.Document(file_path)
            # 读取段落文本
            paragraphs_text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            # 读取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:  # 只添加非空行
                        tables_text.append(" | ".join(row_text))
            # 合并段落和表格内容
            text_content = "\n\n".join(paragraphs_text)
            if tables_text:
                text_content += "\n\n表格内容：\n" + "\n".join(tables_text)
            return text_content
            
        elif suffix in ['.xls', '.xlsx', '.csv']:
            # 处理电子表格
            if suffix == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            # 将 DataFrame 转换为字符串格式
            return df.to_string()
            
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
            
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {str(e)}")
        logger.error(traceback.format_exc())
        raise
        
    finally:
        # 清理临时文件
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.error(f"Error cleaning up temporary file: {str(e)}")

@app.on_event("shutdown")
async def cleanup_all():
    """应用关闭时清理所有临时文件"""
    for session_id in list(app.state.app_state.temp_files.keys()):
        await cleanup_temp_files(session_id)

# 添加清理端点
@app.post("/cleanup")
async def cleanup_session(
    session_id: str = Body(...)
):
    """清理指定会话的临时文件"""
    await cleanup_temp_files(session_id)
    return {"status": "success"}